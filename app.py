import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime

# HARD-CODED PRICING DATA FROM 2026 PRICE SHEET
CRANE_DATA = {
    "Boom Truck": {"rate": 150.00, "ot": 45.00, "sun": 90.00, "in": 0, "out": 0},
    "40 Ton": {"rate": 200.00, "ot": 45.00, "sun": 90.00, "in": 0, "out": 0},
    "60 Ton": {"rate": 240.00, "ot": 45.00, "sun": 90.00, "in": 0, "out": 0},
    "75 Ton": {"rate": 265.00, "ot": 45.00, "sun": 90.00, "in": 0, "out": 0},
    "80 Ton A/T": {"rate": 300.00, "ot": 90.00, "sun": 180.00, "in": 750.00, "out": 750.00},
    "100 Ton A/T": {"rate": 335.00, "ot": 90.00, "sun": 180.00, "in": 1000.00, "out": 1000.00},
    "120 Ton A/T": {"rate": 360.00, "ot": 90.00, "sun": 180.00, "in": 1500.00, "out": 1500.00},
    "165 Ton A/T": {"rate": 400.00, "ot": 90.00, "sun": 180.00, "in": 2500.00, "out": 2500.00},
    "225 Ton A/T": {"rate": 460.00, "ot": 90.00, "sun": 180.00, "in": 3500.00, "out": 3500.00},
    "250 Ton A/T": {"rate": 525.00, "ot": 90.00, "sun": 180.00, "in": 4500.00, "out": 4500.00},
}

st.markdown("<h1 style='text-align: center;'>🏗️ Quote Generator</h1>", unsafe_allow_html=True)

with st.form("quote_input"):
    tonnage = st.selectbox("Crane Tonnage", list(CRANE_DATA.keys()))
    q_date = st.date_input("Date", datetime.today())
    c_name = st.text_input("Company Name")
    c_addr1 = st.text_input("Company Address Line 1")
    c_addr2 = st.text_input("Company Address Line 2")
    c_phone = st.text_input("Company Phone Number")
    contact = st.text_input("Contact Name")
    p_name = st.text_input("Project Name")
    p_addr1 = st.text_input("Project Address Line 1")
    p_addr2 = st.text_input("Project Address Line 2")
    email = st.text_input("Email")
    scope = st.text_area("Scope of Work")
    
    submitted = st.form_submit_button("Generate Quote")

if submitted:
    price_info = CRANE_DATA[tonnage]
    is_heavy = price_info["in"] > 0
    
    # Choose Template based on tonnage
    temp_path = "80 TON AND ABOVE QUOTE TEMPLATE.docx" if is_heavy else "75 TON AND BELOW QUOTE TEMPLATE.docx"
    
    try:
        doc = Document(temp_path)
        
       # This list adds exactly 12 spaces before the OT/Sunday rates for alignment
        replaces = {
            "[Date]": q_date.strftime("%B %d, %Y"),
            "[Company Name]": c_name,
            "[Company Address Line 1]": c_addr1,
            "[Company Address Line 2]": c_addr2,
            "[Address Line 1]": c_addr1,
            "[Address Line 2]": c_addr2,
            "[Company Phone Number]": c_phone,
            "[Company Phone]": c_phone,
            "[Contact Name]": contact,
            "[Project Name]": p_name,
            "[Project Address Line 1]": p_addr1,
            "[Project Address Line 2]": p_addr2,
            "[Email]": email,
            "[Scope of Work]": scope,
            "[scope of work]": scope,
            "[75 Ton]": tonnage,
            "[80 Ton]": tonnage,
            "$000.00 Per Hour": f"${price_info['rate']:.2f} Per Hour",
            
            # THE 12-SPACE ALIGNMENT FIX:
            "$45.00 Per Hour": f"{' ' * 12}${price_info['ot']:.2f} Per Hour",
            "$90.00 Per Hour": f"{' ' * 12}${price_info['sun']:.2f} Per Hour" if not is_heavy else f"{' ' * 12}${price_info['ot']:.2f} Per Hour",
            "$180.00 Per Hour": f"{' ' * 12}${price_info['sun']:.2f} Per Hour",
        }

        # IMPROVED SEARCH: Checks every paragraph and every table cell
        def apply_replacements(container):
            for item in container:
                for key, value in replaces.items():
                    if key in item.text:
                        # This clears the weird hidden Word formatting and forces the text in
                        item.text = item.text.replace(key, value)
                
                # Special logic for Mobilization fees in the 80+ ton template
                if is_heavy:
                    if "Mobilization In" in item.text:
                        item.text = item.text.replace("$000.00", f"${price_info['in']:,.2f}")
                    if "Mobilization Out" in item.text:
                        item.text = item.text.replace("$000.00", f"${price_info['out']:,.2f}")

        # Run the fix on the main text
        apply_replacements(doc.paragraphs)
        
        # Run the fix on any tables (some templates put addresses in tables)
        for table in doc.tables:
            for row in table.rows:
                apply_replacements(row.cells)

        output = BytesIO()
        doc.save(output)
        st.success("Quote Ready!")
        st.download_button("⬇️ Download .docx Quote", output.getvalue(), f"Quote_{c_name}.docx")
    except Exception as e:
        st.error(f"Error: {e}")
